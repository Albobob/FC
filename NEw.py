from openpyxl import load_workbook
import dict
from itertools import groupby
from docx import Document

nz = [int(i) for i in input('Введите номера нозологий чере пробел  ').split()]
# nz = []
# for i in range(100):
#     nz.append(i+1)
district = input('Введите округ  ')
# region = input('Введите регион  ')
region_name = input('Введите регион')
last_year = input('Введите последний год из Формы 2  ')
Yanvar = input('Введите  <январь – февраль> ')
wb = load_workbook(filename='SMP.xlsm')


# last_year = 2020


def decode(wb, sheet, coll, rows):
    sheet = wb[f'{sheet}']
    value_cell = sheet[f'{coll}{rows}'].value
    return value_cell


def record(work_book, sheet, coll, rows, value):
    sheet = work_book[f'{sheet}']
    sheet[f'{coll}{rows}'] = float(value)


def record_one(work_book, sheet, coll, rows, value):
    sheet = work_book[f'{sheet}']
    sheet[f'{coll}{rows}'] = float(value)

rows = 1
key_dict = 1

while rows <= 545 or key_dict <= len(dict.name_two_unit):
    name_form_two = decode(wb, 'Unit', 'B', rows)
    ls_name = dict.name_two_unit.get(key_dict)
    ls_name.append(name_form_two)

    rows += 5
    key_dict += 1


def name_form_two(number_nz):
    name_nz = dict.name_two_unit.get(number_nz)
    name = name_nz[0]

    return name


rows_one_form = 1
key_dict_one = 1

while rows_one_form <= 371 or key_dict_one <= len(dict.name_one_unit):
    name_form_one = decode(wb, 'Unit1', 'B', rows_one_form)
    ls_name_one = dict.name_one_unit.get(key_dict_one)
    ls_name_one.append(name_form_one)

    key_dict_one += 1
    rows_one_form += 5


def name_form_one(number_nz):
    name_nz = dict.name_one_unit.get(number_nz)
    name = name_nz[0]

    return name


for i_one in range(len(dict.name_one_unit)):
    name_one = name_form_one(i_one + 1)
    for i in range(len(dict.name_two_unit)):
        name_two = name_form_two(i + 1)

        if name_two == name_one:
            # print(name_two, i + 1 )
            # print(f'{i + 1}:"{i_one + 1}"', end=',')
            # print(name_one, i_one + 1)

            dict.key_num_two.append(i + 1)
            dict.value_num_one.append(i_one + 1)

# print(len(dict.key_num_two))
# print(len(dict.value_num_one))

for i in range(len(dict.value_num_one)):
    dict.general[f'{int(dict.key_num_two[i])}'] = int(dict.value_num_one[i])

sheet_val = wb['Unit']
sheet_one = wb['Unit1']

cols_cde = ['c', 'd', 'e', 'f', 'g', 'h', 'i', 'j', 'k', 'l']
cols_cde_one = ['c', 'd', 'e', 'f', 'g', 'h', 'i', 'j', 'k', 'l', 'm']

for i in range(108):
    for i_coll in cols_cde:
        ru = sheet_val[f'{i_coll}{dict.number_rows_rus[i]}'].value

        ls_mrb_ru = dict.nz_list_dict_rf.get(i + 1)
        ls_mrb_ru.append(ru)

# Заполняем заболеваемость в District
for i in range(108):
    for i_coll in cols_cde:
        ds = sheet_val[f'{i_coll}{dict.number_rows_dis[i]}'].value
        ls_mrb_ru = dict.nz_list_dict_ds.get(i + 1)
        ls_mrb_ru.append(ds)

# Заполняем заболеваемость в Region
for i in range(109):
    for i_coll in cols_cde:
        rg = sheet_val[f'{i_coll}{dict.number_rows_reg[i]}'].value
        ls_mrb_ru = dict.nz_list_dict_rg.get(i + 1)
        ls_mrb_ru.append(rg)


def mrb_rf(number_nz):
    mrb = dict.nz_list_dict_rf.get(number_nz)
    return mrb


def mrb_ds(number_nz):
    mrb = dict.nz_list_dict_ds.get(number_nz)
    return mrb


def mrb_rg(number_nz):
    mrb = dict.nz_list_dict_rg.get(number_nz)
    return mrb


# ФОРМА 1  //*/*//*//*  ФОРМА 1  //*/*//*//* ФОРМА 1  //*/*//*//*  ФОРМА 1  //*/*//*//* ФОРМА 1  //*/*//*//*  ФОРМА 1

for i in range(108):
    for i_coll in cols_cde_one:
        ru = sheet_one[f'{i_coll}{dict.number_rows_rus_one[i]}'].value

        ls_mrb_ru = dict.nz_list_dict_rf_one.get(i + 1)
        ls_mrb_ru.append(ru)

# Заполняем заболеваемость в District
for i in range(108):
    for i_coll in cols_cde_one:
        ds = sheet_one[f'{i_coll}{dict.number_rows_dis_one[i]}'].value
        ls_mrb_ru = dict.nz_list_dict_ds_one.get(i + 1)
        ls_mrb_ru.append(ds)

# Заполняем заболеваемость в Region
for i in range(109):
    for i_coll in cols_cde_one:
        rg = sheet_one[f'{i_coll}{dict.number_rows_reg_one[i]}'].value
        ls_mrb_ru = dict.nz_list_dict_rg_one.get(i + 1)
        ls_mrb_ru.append(rg)


def one_mrb_rf(number_nz):
    mrb = dict.nz_list_dict_rf_one.get(number_nz)
    return mrb


def one_mrb_ds(number_nz):
    mrb = dict.nz_list_dict_ds_one.get(number_nz)
    return mrb


def one_mrb_rg(number_nz):
    mrb = dict.nz_list_dict_rg_one.get(number_nz)
    return mrb


# print(len(mrb_rg(56)))
# print(len(one_mrb_rg(29)))
#
# print(dict.general)

# for i in dict.general:
#     print(type(i))

def reg_form_one_and_two(number_nz):
    # mrb_form_two = mrb_rg(number_nz)
    mrb_form_one = []

    for i in dict.general:
        if number_nz == int(i):
            value = dict.general.get(str(number_nz))
            # print(value)

            for i in one_mrb_rg(int(value)):
                mrb_form_one.append(i)

    return mrb_form_one


def rus_form_one_and_two(number_nz):
    mrb_form_one = []

    for i in dict.general:
        if number_nz == int(i):
            value = dict.general.get(str(number_nz))
            # print(value)

            for i in one_mrb_rf(int(value)):
                mrb_form_one.append(i)

    return mrb_form_one


def dis_form_one_and_two(number_nz):
    mrb_form_one = []

    for i in dict.general:
        if number_nz == int(i):
            value = dict.general.get(str(number_nz))
            # print(value)

            for i in one_mrb_ds(int(value)):
                mrb_form_one.append(i)

    return mrb_form_one


# print(reg_form_one_and_two(2))

# СМП2 */*/*/* СМП2 *** СМП */*/*/* СМП *** СМП */*/*/* СМП *** СМП */*/*/* СМП2 *** СМП */*/*/* СМП

def smp(number_nz):
    smp_rg_copy = dict.nz_list_dict_rg[number_nz]  # num --номер нозологии из  'nz'
    smp_rg = smp_rg_copy.copy()
    # print(smp_rg)
    smp_rg.pop()  # Удаляем последнее значение (заболеваемость за последний год )
    smp_list = []  # Список в который буду записывать все значения не равные нулю

    # Перебираю заболеваемость в регионе (без последнего года)
    for element in smp_rg:
        if element != 0:
            smp_list.append(element)  # Добавляю значения не равные нулю в список 'smp_list'

    # Нахожу максимальное и минимальное значение
    if sum(smp_list) == 0:
        return "eror"
    elif len(smp_list) < 3:
        return "eror"
    else:

        maximum = max(smp_list)

        idx_max = smp_list.index(maximum)  # Нахожу индекс максимального значения
        smp_list.pop(idx_max)  # Удаляю max значение по 'idx'
        minimum = min(smp_list)  # Нахожу индекс минимального значения
        idx_min = smp_list.index(minimum)  # Удаляю min значение по 'idx'
        smp_list.pop(idx_min)  # Удаляю min значение по 'idx'

        # Нахожу одинаковые рядомстоящие значения
        smp_ls = [el for el, _ in groupby(smp_list)]

        sm = sum(smp_ls)

        if sm != 0:
            smp_value = sum(smp_ls) / len(smp_ls)
            smp_result = "%.2f" % smp_value

            return smp_result
        else:
            return 'EROR'


# СМП1________________________________________________________СМП1
def smp_one(number_nz):
    # print(number_nz)
    smp_rg_copy = dict.nz_list_dict_rg_one[number_nz]  # num --номер нозологии из  'nz'
    smp_rg = smp_rg_copy.copy()
    # print(smp_rg)
    smp_rg.pop()  # Удаляем последнее значение (заболеваемость за последний год )
    smp_list = []  # Список в который буду записывать все значения не равные нулю

    # Перебираю заболеваемость в регионе (без последнего года)
    for element in smp_rg:
        if element != 0:
            smp_list.append(element)  # Добавляю значения не равные нулю в список 'smp_list'

    # Нахожу максимальное и минимальное значение
    if sum(smp_list) == 0:
        return "eror"
    elif len(smp_list) < 3:
        return "eror"
    else:

        maximum = max(smp_list)

        idx_max = smp_list.index(maximum)  # Нахожу индекс максимального значения
        smp_list.pop(idx_max)  # Удаляю max значение по 'idx'
        # print(smp_list)
        minimum = min(smp_list)  # Нахожу индекс минимального значения
        idx_min = smp_list.index(minimum)  # Удаляю min значение по 'idx'
        smp_list.pop(idx_min)  # Удаляю min значение по 'idx'

        # Нахожу одинаковые рядомстоящие значения
        smp_ls = [el for el, _ in groupby(smp_list)]

        sm = sum(smp_ls)

        if sm != 0:
            smp_value = sum(smp_ls) / len(smp_ls)
            smp_result = "%.2f" % smp_value

            return smp_result
        else:
            return 'EROR'


# print(smp(6))
# print(smp_one(60))

'''

def smp_one(number_nz):

    smp_rg = dict.nz_list_dict_rg_one[number_nz]  # number_nz --номер нозологии из  'nz'
        # print(smp_rg)
    smp_rg.pop()  # Удаляем последнее значение (заболеваемость за последний год )
    smp_list = []  # Список в который буду записывать все значения не равные нулю

        # Перебираю заболеваемость в регионе (без последнего года)
    for element in smp_rg:
        if element != 0:
            smp_list.append(element)  # Добавляю значения не равные нулю в список 'smp_list'

        # Нахожу максимальное и минимальное значение
    if sum(smp_list) != 0:

        maximum = max(smp_list)

        idx_max = smp_list.index(maximum)  # Нахожу индекс максимального значения
        smp_list.pop(idx_max)  # Удаляю max значение по 'idx'
        print(smp_list)
        if  sum(smp_list)!= 0:

            print(smp_list)
            minimum = min(smp_list)  # Нахожу индекс минимального значения
                # print(minimum)
                # print(maximum)
            idx_min = smp_list.index(minimum)  # Удаляю min значение по 'idx'
            smp_list.pop(idx_min)  # Удаляю min значение по 'idx'

            # Нахожу одинаковые рядомстоящие значения
        smp_ls = [el for el, _ in groupby(smp_list)]
        sm = sum(smp_ls)
            # print(smp_ls)

        if sm != 0:
            smp_value = sum(smp_ls) / len(smp_ls)
            smp_one = "%.2f" % smp_value

            #return '***EROR!!! (Наверное смп = 0)***'

'''''
# СМП */*/*/* СМП *** СМП */*/*/* СМП *** СМП */*/*/* СМП *** СМП */*/*/* СМП *** СМП */*/*/* СМП
# print()
up_down = {1: 'ниже', 2: 'на уровне', 3: 'выше'}
pretext = {1: 'в', 2: 'на', 3: ' '}
time_percent = {1: 'раза', 2: '%.', 3: ''}

new_list_reg = []
new_list_district = []
new_list_rf = []
new_list_smp = []

for number_nz in nz:
    rus = mrb_rf(number_nz)
    last_value_rus = rus[len(rus) - 1]
    new_list_rf.append(last_value_rus)

    dis = mrb_ds(number_nz)
    last_value_ds = dis[len(dis) - 1]
    new_list_district.append(last_value_ds)

    reg = mrb_rg(number_nz)
    last_value_reg = reg[len(reg) - 1]
    new_list_reg.append(last_value_reg)

documents = Document()

# obj_styles = documents.styles
# obj_charstyle = obj_styles.add_style('CommentsStyle')
# obj_font = obj_charstyle.font
# # obj_font.size = Pt(14)
# obj_font.name = 'Times New Roman'

# paragraph.add_run(any string variable, style = 'CommentsStyle')
p = documents.add_paragraph()  # Создаем параграф
p.add_run('МАТЕРИАЛЫ').bold = True  # Записываем в этот пораграф
p.add_run(f' о деятельности Управления Роспотребнадзора по {region_name}').bold = True
# print(region)
p.add_run(f' и ФБУЗ «Центр гигиены и эпидемиологии в {region_name}»').bold = True
p = documents.add_paragraph()  # Создаем параграф
p.add_run(f'Инфекционная и паразитарная заболеваемость в {int(last_year) - 10} - {last_year}гг.').bold = True
p = documents.add_paragraph('')  # Создаем параграф
p.add_run(f'В {last_year} году для некоторых инфекционных болезней отмечается подъем заболеваемости, ')
p.add_run('что может быть вызвано как ухудшением эпидемиологической ситуации, ')
p.add_run('так и характерными проявлениями эпидемического процесса отдельных ')
p.add_run('инфекций или улучшением качества диагностики среди населения. ')
p.add_run('Так, в многолетней динамике отмечается рост заболеваемости: ')

smp_ls = []
smp_ls_one = []
for i in range(len(nz)):
    value_smp = smp(nz[i])
    smp_ls.append(value_smp)

    # for i in range(len(nz)):
    print('*******************')
    vl = nz[i]
    name_nz = dict.name_two_unit.get(vl)
    name = name_nz[0]
    region_value = mrb_rg(nz[i])[len(mrb_rg(nz[i])) - 1]
    district_value = mrb_ds(nz[i])[len(mrb_ds(nz[i])) - 1]
    russia_value = mrb_rf(nz[i])[len(mrb_rf(nz[i])) - 1]

    up_down_list = []
    pretext_list = []
    time_percent_list = []

    time_percent_value_list = []  # Если < 1.5 раз то %, если > 1.5 то разы

    percent = (new_list_reg[i] * 100) / new_list_rf[i] - 100  # Узнаем в процентах
    time = 0.0  # узнаем во сколько раз больше
    if new_list_reg[i] > new_list_rf[i]:
        time = new_list_reg[i] / new_list_rf[i]
        up_down_list.append(up_down.get(3))  # если регион больше РФ то  показатель заболеваемости (up_down) = ВЫШЕ
    else:
        time = new_list_rf[i] / new_list_reg[i]
        up_down_list.append(up_down.get(1))  # если РФ больше региона то  показатель заболеваемости (up_down) = НИЖЕ

    if abs(time) > 1.5:
        pretext_list.append(pretext.get(1))  # в
        time_value = "%.2f" % abs(time)
        time_percent_value_list.append(time_value)
        time_percent_list.append(time_percent.get(1))

    if abs(time) <= 1.5:
        if 0 < abs(percent) < 12:  # ПРОЦЕНТЫ на 'уровне'
            up_down_list.clear()
            up_down_list.append(up_down.get(2))
            pretext_list.append(pretext.get(3))
            time_percent_value_list.append('')
            time_percent_list.append(time_percent.get(3))
        else:
            pretext_list.append(pretext.get(2))
            percent_value = "%.1f" % abs(percent)
            time_percent_value_list.append(percent_value)
            time_percent_list.append(time_percent.get(2))

    upd = up_down_list[0]  # выше / на уровне / ниже
    pt_l = pretext_list[0]  # в / на
    ttp_value = time_percent_value_list[0]  # значение
    tp_l = time_percent_list[0]  # в % или раз
    rs = f'Рис.{i + 1} '

    p = documents.add_paragraph()  # Создаем параграф
    p.add_run(
        f'{name} – ').bold = True
    p.add_run(
        f' {region_value} на 100 тыс. населения при среднемноголетней заболеваемости {smp_ls[i]} (Рис.{i + 1}).'
        f'Показатель по субъекту в {last_year} году')
    p.add_run(f' {upd} ').bold = True
    p.add_run(f'показателя по Российской Федерации ({russia_value} на 100 тыс. населения)')
    p.add_run(f' {pt_l} {time_percent_value_list[0]} {tp_l}. ').bold = True
    p.add_run(
        f' Заболеваемость {name} в {district} в {last_year}г. составила {district_value} на 100 тыс. населения.\n')
    p.add_run(rs)

    for i in dict.general:

        if int(vl) == int(i):
            one_rg = reg_form_one_and_two(int(vl))
            one_rg_vl = one_rg[len(one_rg) - 1]

            one_rf = rus_form_one_and_two(int(vl))
            one_rf_vl = one_rf[len(one_rf) - 1]

            one_ds = dis_form_one_and_two(int(vl))
            one_ds_vl = one_ds[len(one_ds) - 1]

            num_smp_one = dict.general.get(f'{vl}')
            # print(num_smp_one)
            # print(vl)
            value = smp_one(num_smp_one)
            # print(value)
            smp_ls_one.append(value)
            # print(smp_ls_one)
            # выше ниже форма111111111111111111111111111111111111111111111111111111111111111
            up_down_list_one = []
            pretext_list_one = []
            time_percent_list_one = []

            time_percent_value_list_one = []  # Если < 1.5 раз то %, если > 1.5 то разы

            percent = (one_rg_vl * 100) / one_rf_vl - 100  # Узнаем в процентах
            time = 0.0  # узнаем во сколько раз больше
            if one_rg_vl != 0:
                if one_rg_vl > one_rf_vl:
                    time = one_rg_vl / one_rf_vl
                    up_down_list_one.append(
                        up_down.get(3))  # если регион больше РФ то  показатель заболеваемости (up_down) = ВЫШЕ
                else:
                    time = one_rf_vl / one_rg_vl
                    up_down_list_one.append(
                        up_down.get(1))  # если РФ больше региона то  показатель заболеваемости (up_down) = НИЖЕ

                if abs(time) > 1.5:
                    pretext_list_one.append(pretext.get(1))  # в
                    time_value = "%.2f" % abs(time)
                    time_percent_value_list_one.append(time_value)
                    time_percent_list_one.append(time_percent.get(1))

                if abs(time) <= 1.5:
                    if 0 < abs(percent) < 12:  # ПРОЦЕНТЫ на 'уровне'
                        up_down_list_one.clear()
                        up_down_list_one.append(up_down.get(2))
                        pretext_list_one.append(pretext.get(3))
                        time_percent_value_list_one.append('')
                        time_percent_list_one.append(time_percent.get(3))
                    else:
                        pretext_list_one.append(pretext.get(2))
                        percent_value = "%.1f" % abs(percent)
                        time_percent_value_list_one.append(percent_value)
                        time_percent_list_one.append(time_percent.get(2))
            if one_rg_vl != 0:
                upd_one = up_down_list_one[0]  # выше / на уровне / ниже
                pt_l_one = pretext_list_one[0]  # в / на
                ttp_value_one = time_percent_value_list_one[0]  # значение
                tp_l_one = time_percent_list_one[0]  # в % или раз
                probel = ' '
            else:
                upd_one = ''
                pt_l_one = ''
                ttp_value_one = ''
                tp_l_one = ''
                probel = ''

            if one_rg_vl == 0:
                p.add_run(
                    f'За {Yanvar} {int(last_year) + 1}г.').bold = True
                p.add_run(
                    f' не выявлено ни одного случая заболевания в')
                p.add_run(f' {region_name}')
                p.add_run(
                    f' (СМП - {value}), показатель по Российской Федерации ({one_rf_vl}){probel}{pt_l_one}{probel}{ttp_value_one}{probel}{tp_l_one}.'
                    f'Показатель по {district} составил {one_ds_vl} на 100 тыс. населения. \n')
                p.add_run(rs)
            else:
                p.add_run(
                    f'За {Yanvar} {int(last_year) + 1}г.').bold = True
                p.add_run(f' Показатель заболеваемости составил ')
                p.add_run(f'{one_rg_vl}').bold = True
                p.add_run(
                    f'(СМП - {value}), что ')
                p.add_run(f'{upd_one} ').bold = True
                p.add_run(
                    f'показателя по Российской Федерации ({one_rf_vl}){probel}{pt_l_one}{probel}{ttp_value_one}{probel}{tp_l_one}.'
                    f'Показатель по {district} составил {one_ds_vl} на 100 тыс. населения. \n')
                p.add_run(rs)

for i in nz:
    w = int(i) * 5 - 2
    value = smp(i)
    record(wb, 'Unit', 'P', w,  float(value))


    for io in dict.general:
        if int(i) == int(io):
            print(i)
            x = dict.general.get(str(i))
            value_one = smp_one(int(x))
            w_one = x * 5 - 2
            print(x)
            record_one(wb, 'Unit1', 'P', w_one, float(value))



documents.save(f'{region_name}+{nz}.docx')  # ПОСЛЕДНЯЯ СТРОЧКА В КОДЕ!!!
wb.save(f'!ТАБЛИЦЫ{region_name}.xlsx')
