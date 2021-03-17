import dict
from openpyxl import load_workbook
from itertools import groupby
from docx import Document

# MENU
file_name = 'SMP.xlsm'
sheet = 'Unit'
nz = [int(i) for i in input('Введите номера нозологий чере пробел  ').split()]
district = input('Введите округ  ')
region = input('Введите регион  ')
last_year = input('Введите последний год')
first_year = input('Введите первый год')

# Работа c EXEL***...***Работа c EXEL***...***Работа c EXEL***...***Работа c EXEL***...***Работа c EXEL***...***EXEL////
wb = load_workbook(filename=f'{file_name}')
sheet_val = wb[f'{sheet}']

cols_cde = ['c', 'd', 'e', 'f', 'g', 'h', 'i', 'j', 'k', 'l']

for i in range(108):
    for i_coll in cols_cde:
        ru = sheet_val[f'{i_coll}{dict.number_rows_rus[i]}'].value

        ls_mrb_ru = dict.nz_list_dict_rf.get(i + 1)
        ls_mrb_ru.append(ru)

# Заполняем заболеваемость в District
for i in range(108):
    for i_coll in cols_cde:
        ds = sheet_val[f'{i_coll}{dict.number_rows_dis[i]}'].value
        ls_mrb_ru = dict.nz_list_dict_ds.get(i + 1)
        ls_mrb_ru.append(ds)

# Заполняем заболеваемость в Region
for i in range(109):
    for i_coll in cols_cde:
        rg = sheet_val[f'{i_coll}{dict.number_rows_reg[i]}'].value
        ls_mrb_ru = dict.nz_list_dict_rg.get(i + 1)
        ls_mrb_ru.append(rg)

# Рисуем табличку в tbl.xlsx

wb_tbl = load_workbook(filename='tbl.xlsx')
sheet_rec = wb_tbl['Лист1']
unt_tbl = len(nz) * 5


def mrb_rf(number_nz):
    mrb = dict.nz_list_dict_rf.get(number_nz)
    return mrb


def mrb_ds(number_nz):
    mrb = dict.nz_list_dict_ds.get(number_nz)
    return mrb


def mrb_rg(number_nz):
    mrb = dict.nz_list_dict_rg.get(number_nz)
    return mrb


# def ru(coll, rows, number_nz, value):
#     sheet_rec[f'{coll}{rows}'] = f'{mrb_rf(number_nz)[value]}'


rows_nz = 2
# rows_year = 3
rows_rf_ls = []
rows_ds_ls = []
rows_rg_ls = []
rows_year_ls = []
rows_nz_ls = []

while rows_nz <= unt_tbl:
    rows_nz_ls.append(rows_nz + 0)
    rows_year_ls.append(rows_nz + 1)
    rows_rf_ls.append(rows_nz + 2)
    rows_ds_ls.append(rows_nz + 3)
    rows_rg_ls.append(rows_nz + 4)
    # rows_rg_ls.append(rows_year + 4)

    rows_nz += 5

print(rows_year_ls)


def ru(nz, rows_rf_ls, cols_cde):
    for i in range(len(nz)):
        # print(i)
        string = (rows_rf_ls[i])
        ls = (mrb_rf(nz[i]))

        for i_coll in range(10):
            coll = cols_cde[i_coll]
            # print(coll)
            # print(ls[i_coll])
            sheet_rec[f'{coll}{string}'] = f'{ls[i_coll]}'


def ds(nz, rows_ds_ls, cols_cde):
    for i in range(len(nz)):
        # print(i)
        string = (rows_ds_ls[i])
        ls = (mrb_ds(nz[i]))

        for i_coll in range(10):
            coll = cols_cde[i_coll]
            # print(coll)
            # print(ls[i_coll])
            sheet_rec[f'{coll}{string}'] = f'{ls[i_coll]}'


def rg(nz, rows_rg_ls, cols_cde):
    for i in range(len(nz)):
        # print(i)
        string = (rows_rg_ls[i])
        ls = (mrb_rg(nz[i]))

        for i_coll in range(10):
            coll = cols_cde[i_coll]
            # print(coll)
            # print(ls[i_coll])
            sheet_rec[f'{coll}{string}'] = f'{ls[i_coll]}'


year = int(first_year)
year_ls = []

while year <= int(last_year):
    year_ls.append(year)

    year += 1
# ЗАПИСЫВАЕМ ГОД
for i in range(len(nz)):
    string = (rows_year_ls[i])
    for i_coll in range(10):
        coll = cols_cde[i_coll]
        y_v = year_ls[i_coll]
        #     print(coll)
        # print(y_v)
        # print(i)
        # print(string)
        sheet_rec[f'{coll}{string}'] = f'{y_v}'

# ЗАПИСЫВАЕМ НАЗВАНИЕ НОЗОЛОГИИ
for i in range(len(nz)):
    string = (rows_nz_ls[i])
    y_v = dict.nz_dict.get(nz[i])
    print(y_v)
    sheet_rec[f'B{string}'] = f'{y_v}'
    sheet_rec[f'B{string + 1 + 1}'] = f'Российская Федирация'
    sheet_rec[f'B{string + 2 + 1}'] = f'{district}'
    sheet_rec[f'B{string + 3 + 1}'] = f'{region}'
# ЗАПИСЫВАЕМ ЗАБОЛЕВАЕМОСТЬ
ru(nz, rows_rf_ls, cols_cde)
ds(nz, rows_ds_ls, cols_cde)
rg(nz, rows_rg_ls, cols_cde)

wb_tbl.save('demo_tbl.xlsx')


# РАБОТА К ВОРДУ


def smp(number_nz):
    smp_rg = dict.nz_list_dict_rg[number_nz]  # num --номер нозологии из  'nz'
    # print(smp_rg)
    smp_rg.pop()  # Удаляем последнее значение (заболеваемость за последний год )
    smp_list = []  # Список в который буду записывать все значения не равные нулю

    # Перебираю заболеваемость в регионе (без последнего года)
    for element in smp_rg:
        if element != 0:
            smp_list.append(element)  # Добавляю значения не равные нулю в список 'smp_list'

    # Нахожу максимальное и минимальное значение
    maximum = max(smp_list)
    idx_max = smp_list.index(maximum)  # Нахожу индекс максимального значения
    smp_list.pop(idx_max)  # Удаляю max значение по 'idx'
    minimum = min(smp_list)  # Нахожу индекс минимального значения
    idx_min = smp_list.index(minimum)  # Удаляю min значение по 'idx'
    smp_list.pop(idx_min)  # Удаляю min значение по 'idx'

    # Нахожу одинаковые рядомстоящие значения
    smp_ls = [el for el, _ in groupby(smp_list)]

    smp_value = sum(smp_ls) / len(smp_ls)
    smp = "%.2f" % smp_value

    return smp


new_list_reg = []
new_list_district = []
new_list_rf = []
new_list_smp = []

for number_nz in nz:
    rus = mrb_rf(number_nz)
    last_value_rus = rus[len(rus) - 1]
    new_list_rf.append(last_value_rus)

    dis = mrb_ds(number_nz)
    last_value_ds = dis[len(dis) - 1]
    new_list_district.append(last_value_ds)

    reg = mrb_rg(number_nz)
    last_value_reg = reg[len(reg) - 1]
    new_list_reg.append(last_value_reg)

    s = smp(number_nz)
    new_list_smp.append(s)

print(new_list_rf)
print(new_list_reg)
print(new_list_smp)

# ТЕКСТОВКА...***...ТЕКСТОВКА...***...ТЕКСТОВКА...***...ТЕКСТОВКА...***...ТЕКСТОВКА...***...ТЕКСТОВКА...***...ТЕКСТОВКА
documents = Document()

up_down = {1: 'ниже', 2: 'на уровне', 3: 'выше'}
pretext = {1: 'в', 2: 'на', 3: ' '}
time_percent = {1: 'раза.', 2: '%.', 3: ''}

p = documents.add_paragraph()  # Создаем параграф
p.add_run('МАТЕРИАЛЫ').bold = True  # Записываем в этот пораграф
p.add_run(f' о деятельности Управления Роспотребнадзора по {region}').bold = True
p.add_run(f' и ФБУЗ «Центр гигиены и эпидемиологии в {region}»').bold = True
p = documents.add_paragraph()  # Создаем параграф
p.add_run(f'Инфекционная и паразитарная заболеваемость в {first_year} - {last_year}гг.').bold = True
p = documents.add_paragraph('')  # Создаем параграф
p.add_run(f'В {last_year} году для некоторых инфекционных болезней отмечается подъем заболеваемости, ')
p.add_run('что может быть вызвано как ухудшением эпидемиологической ситуации, ')
p.add_run('так и характерными проявлениями эпидемического процесса отдельных ')
p.add_run('инфекций или улучшением качества диагностики среди населения. ')
p.add_run('Так, в многолетней динамике отмечается рост заболеваемости: ')

p = documents.add_paragraph()  # Создаем параграф
p = documents.add_paragraph()  # Создаем параграф
p = documents.add_paragraph()  # Создаем параграф

for i in range(len(nz)):
    p = documents.add_paragraph()  # Создаем параграф
    p.add_run(f'{dict.nz_dict.get(nz[i])} – ').bold = True
    p.add_run(
        f'{new_list_reg[i]} на 100 тыс. населения при среднемноголетней заболеваемости {new_list_smp[i]}. Показатель ')
    p.add_run(f'по субъекту в {last_year} году')

    up_down_list = []
    pretext_list = []
    time_percent_list = []

    time_percent_value_list = []  # Если < 1.5 раз то %, если > 1.5 то разы

    percent = (new_list_reg[i] * 100) / new_list_rf[i] - 100  # Узнаем в процентах
    time = 0.0  # узнаем во сколько раз больше
    if new_list_reg[i] > new_list_rf[i]:
        time = new_list_reg[i] / new_list_rf[i]
        up_down_list.append(up_down.get(3))  # если регион больше РФ то  показатель заболеваемости (up_down) = ВЫШЕ
    else:
        time = new_list_rf[i] / new_list_reg[i]
        up_down_list.append(up_down.get(1))  # если РФ больше региона то  показатель заболеваемости (up_down) = НИЖЕ

    if abs(time) > 1.5:
        pretext_list.append(pretext.get(1))  # в
        time_value = "%.2f" % abs(time)
        time_percent_value_list.append(time_value)
        time_percent_list.append(time_percent.get(1))

    if abs(time) <= 1.5:
        if 0 < abs(percent) < 12:  # ПРОЦЕНТЫ на 'уровне'
            up_down_list.clear()
            up_down_list.append(up_down.get(2))
            pretext_list.append(pretext.get(3))
            time_percent_value_list.append('')
            time_percent_list.append(time_percent.get(3))
        else:
            pretext_list.append(pretext.get(2))
            percent_value = "%.1f" % abs(percent)
            time_percent_value_list.append(percent_value)
            time_percent_list.append(time_percent.get(2))

    upd = up_down_list[0]  # выше / на уровне / ниже
    pt_l = pretext_list[0]  # в / на
    ttp_value = time_percent_value_list[0]  # значение
    tp_l = time_percent_list[0]  # в % или раз

    p.add_run(f' {upd} ').bold = True
    p.add_run(f'показателя по Российской Федерации ({new_list_rf[i]} на 100 тыс. населения) ')
    p.add_run(f'{pt_l} {time_percent_value_list[0]} {tp_l}').bold = True
    p.add_run(
        f'  Заболеваемость {dict.nz_dict.get(nz[i])} в  {district} в {last_year} составила {new_list_district[i]} на 100 тыс. населения')
    p = documents.add_paragraph()  # Создаем параграф
    p.add_run(
        f'Рис.{i + 1} Заболеваемость {dict.nz_dict.get(nz[i])} в {region} в {first_year} - {last_year}гг. (на 100 тыс. населения).')

documents.save(f'{region}+{nz}.docx')  # ПОСЛЕДНЯЯ СТРОЧКА В КОДЕ!!!
